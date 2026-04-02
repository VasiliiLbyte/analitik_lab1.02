"""Commercial Proposal (КП) Word document generator.

Builds .docx files matching the exact format from the PDF examples:
  - Header: "Коммерческое предложение № XX от DD MMMM YYYY г."
  - Executor: full legal details of ООО "АНАЛИТИК.ЛАБ"
  - Customer: data from KP form
  - Table: services grouped by category with sub-totals
  - Protocol fee (3%), NDS (5%), Grand total
  - Total in words via num2words

Uses python-docx directly for full control over cell formatting.
No fragile docxtpl row-loop tags.
"""

from __future__ import annotations

import datetime
import logging
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from num2words import num2words

from bot.config import GENERATED_DIR, get_settings

logger = logging.getLogger(__name__)

_MONTHS_RU = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля",
    5: "мая", 6: "июня", 7: "июля", 8: "августа",
    9: "сентября", 10: "октября", 11: "ноября", 12: "декабря",
}


@dataclass
class KPLineItem:
    number: int
    name: str
    quantity: int
    unit: str
    unit_price: float
    total_price: float


@dataclass
class KPCategoryGroup:
    name: str
    items: list[KPLineItem] = field(default_factory=list)
    subtotal: float = 0.0


@dataclass
class KPData:
    kp_number: str
    kp_date: datetime.date
    customer_name: str
    customer_inn: str
    customer_kpp: str
    customer_address: str
    contact_person: str
    contact_info: str
    sample_location: str = ""
    research_deadline: str = ""
    sample_return: str = ""
    groups: list[KPCategoryGroup] = field(default_factory=list)
    subtotal: float = 0.0
    protocol_fee: float = 0.0
    total_before_nds: float = 0.0
    nds: float = 0.0
    total: float = 0.0
    total_words: str = ""
    total_items: int = 0


def _format_date_ru(d: datetime.date) -> str:
    return f"{d.day} {_MONTHS_RU[d.month]} {d.year} г."


def _amount_in_words(amount: float) -> str:
    """'Тридцать семь тысяч ... рублей XX копеек'."""
    rubles = int(amount)
    kopecks = round((amount - rubles) * 100)
    rubles_text = num2words(rubles, lang="ru").capitalize()
    return f"{rubles_text} {_decline_rubles(rubles)} {kopecks:02d} {_decline_kopecks(kopecks)}"


def _decline_rubles(n: int) -> str:
    last2 = n % 100
    last1 = n % 10
    if 11 <= last2 <= 19:
        return "рублей"
    if last1 == 1:
        return "рубль"
    if 2 <= last1 <= 4:
        return "рубля"
    return "рублей"


def _decline_kopecks(n: int) -> str:
    last2 = n % 100
    last1 = n % 10
    if 11 <= last2 <= 19:
        return "копеек"
    if last1 == 1:
        return "копейка"
    if 2 <= last1 <= 4:
        return "копейки"
    return "копеек"


def _fmt(value: float) -> str:
    return f"{value:,.2f}".replace(",", " ")


def _set_cell(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT,
              size: int = 10) -> None:
    """Write text into a table cell with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.bold = bold


def build_kp_data(
    kp_number: str,
    kp_date: datetime.date,
    customer_name: str,
    customer_inn: str,
    customer_kpp: str,
    customer_address: str,
    contact_person: str,
    contact_info: str,
    cart_items: list,
    sample_location: str = "",
    research_deadline: str = "",
    sample_return: str = "",
) -> KPData:
    """Build KPData from cart items, grouping by category with auto protocol fee."""
    settings = get_settings()

    grouped: dict[str, list] = defaultdict(list)
    for item in cart_items:
        grouped[item.category_name].append(item)

    groups: list[KPCategoryGroup] = []
    global_num = 0
    subtotal = 0.0
    total_items = 0

    for cat_name, items in grouped.items():
        group = KPCategoryGroup(name=cat_name)
        cat_subtotal = 0.0
        for item in items:
            global_num += 1
            total_items += item.quantity
            line_total = round(item.unit_price * item.quantity, 2)
            cat_subtotal += line_total
            group.items.append(KPLineItem(
                number=global_num,
                name=item.service_name,
                quantity=item.quantity,
                unit=item.unit,
                unit_price=item.unit_price,
                total_price=line_total,
            ))
        group.subtotal = round(cat_subtotal, 2)
        subtotal += cat_subtotal
        groups.append(group)

    subtotal = round(subtotal, 2)
    protocol_fee = round(subtotal * settings.protocol_fee_rate / 100, 2)
    total_before_nds = round(subtotal + protocol_fee, 2)
    nds = round(total_before_nds * settings.nds_rate / 100, 2)
    total = round(total_before_nds + nds, 2)

    return KPData(
        kp_number=kp_number,
        kp_date=kp_date,
        customer_name=customer_name,
        customer_inn=customer_inn,
        customer_kpp=customer_kpp,
        customer_address=customer_address,
        contact_person=contact_person,
        contact_info=contact_info,
        sample_location=sample_location,
        research_deadline=research_deadline,
        sample_return=sample_return,
        groups=groups,
        subtotal=subtotal,
        protocol_fee=protocol_fee,
        total_before_nds=total_before_nds,
        nds=nds,
        total=total,
        total_words=_amount_in_words(total),
        total_items=total_items,
    )


def generate_kp(data: KPData) -> Path:
    """Build the KP .docx directly with python-docx and return the file path."""
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    safe_number = data.kp_number.replace("/", "-")
    output_path = GENERATED_DIR / f"KP_{safe_number}_{data.kp_date.isoformat()}.docx"

    settings = get_settings()
    doc = Document()

    # --- Global font ---
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)

    # --- Title ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        f"Коммерческое предложение № {data.kp_number} "
        f"от {_format_date_ru(data.kp_date)}"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    # --- Executor ---
    p = doc.add_paragraph()
    run = p.add_run("Исполнитель:\n")
    run.bold = True
    run.font.name = "Times New Roman"
    p.add_run(
        f'{settings.executor_name}, ИНН {settings.executor_inn}, '
        f'КПП {settings.executor_kpp}, {settings.executor_address}'
    ).font.name = "Times New Roman"

    # --- Customer ---
    p = doc.add_paragraph()
    run = p.add_run("Заказчик:\n")
    run.bold = True
    run.font.name = "Times New Roman"
    p.add_run(
        f'{data.customer_name}, ИНН {data.customer_inn}, '
        f'КПП {data.customer_kpp}, {data.customer_address}'
    ).font.name = "Times New Roman"

    # --- Additional application conditions ---
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Дополнительные условия заявки:")
    run.bold = True
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.add_run(f"Фактическое местоположение объекта: {data.sample_location or 'не указано'}").font.name = "Times New Roman"
    p = doc.add_paragraph()
    p.add_run(f"Сроки проведения исследований: {data.research_deadline or 'не требуется'}").font.name = "Times New Roman"
    p = doc.add_paragraph()
    p.add_run(
        f"Возврат неиспользованной части проб: {data.sample_return or 'не указано'}"
    ).font.name = "Times New Roman"

    # --- Service table ---
    COL_WIDTHS = [Cm(1), Cm(9), Cm(1.5), Cm(1.5), Cm(2.5), Cm(2.5)]
    HEADERS = ["№", "Товары (работы, услуги)", "Кол-во", "Ед.", "Цена", "Сумма"]
    CENTER = WD_ALIGN_PARAGRAPH.CENTER
    RIGHT = WD_ALIGN_PARAGRAPH.RIGHT

    num_cols = len(HEADERS)
    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, (hdr, w) in enumerate(zip(HEADERS, COL_WIDTHS)):
        cell = table.rows[0].cells[i]
        _set_cell(cell, hdr, bold=True, align=CENTER)
        cell.width = w

    # Rows: iterate groups and items
    for group in data.groups:
        # Category header row (merged across all columns)
        row = table.add_row()
        _set_cell(row.cells[0], "")
        _set_cell(row.cells[1], group.name, bold=True)
        _set_cell(row.cells[2], "")
        _set_cell(row.cells[3], "")
        _set_cell(row.cells[4], "")
        _set_cell(row.cells[5], _fmt(group.subtotal), bold=True, align=RIGHT)

        for item in group.items:
            row = table.add_row()
            _set_cell(row.cells[0], str(item.number), align=CENTER)
            _set_cell(row.cells[1], item.name)
            _set_cell(row.cells[2], str(item.quantity), align=CENTER)
            _set_cell(row.cells[3], item.unit, align=CENTER)
            _set_cell(row.cells[4], _fmt(item.unit_price), align=RIGHT)
            _set_cell(row.cells[5], _fmt(item.total_price), align=RIGHT)

    # Protocol fee row
    row = table.add_row()
    _set_cell(row.cells[0], "")
    _set_cell(row.cells[1], "Оформление протоколов (3% от сметной стоимости)", bold=True)
    _set_cell(row.cells[2], "")
    _set_cell(row.cells[3], "")
    _set_cell(row.cells[4], "")
    _set_cell(row.cells[5], _fmt(data.protocol_fee), bold=True, align=RIGHT)

    # --- Totals ---
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run(f"Итого: {_fmt(data.subtotal + data.protocol_fee)} руб.").font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.add_run(f"Кроме того НДС (5%): {_fmt(data.nds)} руб.").font.name = "Times New Roman"

    # Grand total (bold)
    p = doc.add_paragraph()
    run = p.add_run(
        f"Всего наименований {data.total_items}, "
        f"на сумму {_fmt(data.total)} руб."
    )
    run.bold = True
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    run = p.add_run(data.total_words)
    run.font.name = "Times New Roman"

    try:
        doc.save(str(output_path))
        logger.info("KP generated: %s", output_path)
        return output_path
    except Exception:
        logger.exception("Failed to save KP document")
        raise
